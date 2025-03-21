import os
import json
import tempfile
from docx import Document
from transformers import AutoModelForCausalLM, AutoTokenizer
import torch
from dotenv import load_dotenv

load_dotenv()
HF_TOKEN = os.get_env("HF_TOKEN")

# Initialize the locally saved model and tokenizer
def load_model_and_tokenizer():
    """Load the locally saved model and tokenizer."""
    model_path = "aaditey932/llama-resume-optimizer"
    print("Loading model and tokenizer...")
    
    # Load the tokenizer
    tokenizer = AutoTokenizer.from_pretrained(model_path)
    
    # Load the model
    model = AutoModelForCausalLM.from_pretrained(
        model_path,
        torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,  # Use FP16 if GPU is available
        device_map="auto",  # Automatically map the model to available devices (GPU/CPU)
        load_in_4bit=True,  # Load in 4-bit if using quantization
        token = HF_TOKEN
    )
    
    # Enable faster inference
    model.eval()
    return model, tokenizer

# Load the model and tokenizer
model, tokenizer = load_model_and_tokenizer()

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file."""
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def generate_suggestions(resume_text, job_description):
    """Generates ATS-friendly resume optimization suggestions using the locally saved model."""
    prompt = f"""
    You are an AI-powered resume optimization assistant. Your task is to analyze the provided resume against the job description and generate ATS-friendly, structured, and impact-driven improvements.

    ### Requirements:
    - Output should be used programmatically to modify the resume structure.
    - Ensure maximum ATS compatibility:
      - Embed relevant keywords from the job description naturally.
      - Use active, impact-driven language to highlight achievements.
      - Maintain consistent job titles, dates, and formatting.
    - Response must be a structured JSON object.

    ### Output Format:
    Generate a structured JSON object with numbered edits ("edit1", "edit2", etc.):
    - "to_be_edited": The exact resume text needing improvement.
    - "edited": The optimized version with ATS-friendly phrasing.
      - Add missing keywords from the job description.
      - Include impact-driven metrics (e.g., "Increased efficiency by 30%").
      - Improve formatting (e.g., consistent bullet points and dates).

    ### Resume:
    {resume_text}

    ### Job Description:
    {job_description}

    ### Expected JSON Output Example:
    {{
        "edit1": {{
            "to_be_edited": "Led a team of software engineers to develop a recommendation engine for e-commerce.",
            "edited": "Managed a cross-functional team of 5 engineers to build a recommendation engine, increasing conversion rates by 20%."
        }},
        "edit2": {{
            "to_be_edited": "Worked with Python, SQL, and machine learning models to optimize product recommendations.",
            "edited": "Developed and deployed machine learning models using Python and SQL, improving recommendation accuracy by 15%."
        }}
    }}
    Give only the JSON output.
    """

    # Tokenize the input
    inputs = tokenizer(
        prompt,
        return_tensors="pt",
        truncation=True,
        max_length=4096,  # Adjust based on your model's max sequence length
    ).to(model.device)

    # Generate the output
    outputs = model.generate(
        input_ids=inputs["input_ids"],
        attention_mask=inputs["attention_mask"],
        max_new_tokens=1024,  # Adjust based on the desired output length
        use_cache=True,
    )

    # Decode the output
    response = tokenizer.decode(outputs[0], skip_special_tokens=True)
    print("Generated response:", response)

    # Parse the response into a JSON object
    try:
        edits = json.loads(response)
        return edits
    except json.JSONDecodeError:
        print("Failed to parse the model's response as JSON.")
        return None

def replace_text_preserve_formatting(doc, old_text, new_text):
    """
    Replace `old_text` with `new_text` in the document while preserving all formatting.
    Handles cases where `old_text` spans multiple runs.
    """
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            # Initialize variables to track runs and their text
            runs = paragraph.runs
            text = ''.join(run.text for run in runs)
            
            if old_text in text:
                # Find the start and end positions of the old text
                start = text.find(old_text)
                end = start + len(old_text)
                
                # Clear the paragraph and rebuild it
                paragraph.clear()
                
                # Add text before the old text
                if start > 0:
                    paragraph.add_run(text[:start])
                
                # Add the new text with formatting from the first run of the old text
                new_run = paragraph.add_run(new_text)
                # Find the first run that contains part of the old text
                for run in runs:
                    if run.text in old_text:
                        # Copy formatting from this run
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size
                        new_run.font.color.rgb = run.font.color.rgb
                        break
                
                # Add text after the old text
                if end < len(text):
                    paragraph.add_run(text[end:])

def replace_text_in_document(resume_path, output_path, edits):
    """Replaces text in a DOCX file based on provided JSON edits."""
    doc = Document(resume_path)

    for edit_key, edit_value in edits.items():
        old_text = edit_value["to_be_edited"]
        new_text = edit_value["edited"]
        replace_text_preserve_formatting(doc, old_text, new_text)

    doc.save(output_path)

    print(f"✅ Edited DOCX saved to: {output_path}")

    return output_path

def process_resume(resume_file, job_description):
    """Main function to handle resume optimization workflow."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_input:
        temp_input.write(resume_file.read())
        input_path = temp_input.name

    output_path = input_path.replace(".docx", "_edited.docx")

    # Extract text
    resume_text = extract_text_from_docx(input_path)

    # Generate suggestions
    edits = generate_suggestions(resume_text, job_description)

    if edits:
        edited_resume = replace_text_in_document(input_path, output_path, edits)
        os.remove(input_path)  # Clean up temp files
        return edited_resume
    else:
        return None