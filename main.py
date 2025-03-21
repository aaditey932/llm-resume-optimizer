import streamlit as st
import pandas as pd
import re
import pickle
import nltk
from nltk.tokenize import word_tokenize
import PyPDF2
from docx import Document
import numpy as np
import os
import sys
import warnings
from scripts.deep.eval_deep_model import evaluate_resume

# Silence warnings
warnings.filterwarnings("ignore", category=RuntimeWarning)
warnings.filterwarnings("ignore", category=UserWarning)

# Download necessary NLTK data
nltk.download('punkt')
nltk.download('punkt_tab')
nltk.download('wordnet')
nltk.download('stopwords')
nltk.download('averaged_perceptron_tagger')


# Import torch conditionally to avoid Streamlit watcher issues
try:
    import torch
except RuntimeError:
    pass

# Add scripts directory to path for imports
sys.path.append(os.path.join(os.path.dirname(os.path.abspath(__file__)), "scripts", "deep"))


from scripts.deep.trial_gpt_resume_optimizer import process_resume

# Import traditional functions from the combined script
from scripts.traditional.traditional_combined import (
    remove_markup, preprocess_text, jaccard_similarity, ngram_overlap,
    SentenceTransformer, TfidfVectorizer, CountVectorizer, cosine_similarity
)

# Load Pretrained Job Keywords
@st.cache_resource
def load_keywords():
    try:
        with open("./models/category_keywords.pkl", "rb") as file:
            category_keywords = pickle.load(file)
        return category_keywords
    except FileNotFoundError:
        st.warning("Job category keywords file not found. Basic keyword matching will be limited.")
        return {}

# Load BERT model if available
@st.cache_resource
def load_bert_model():
    try:
        # Using a try-except block with a more specific import
        from sentence_transformers import SentenceTransformer
        model = SentenceTransformer('all-mpnet-base-v2')
        return model
    except Exception as e:
        st.warning(f"Could not load BERT model: {str(e)}")
        return None

# Load trained models
@st.cache_resource
def load_models():
    try:
        with open("./models/scaler_features.pkl", "rb") as f:
            scaler_features = pickle.load(f)
        with open("./models/scaler_gmm.pkl", "rb") as f:
            scaler_gmm = pickle.load(f)
        with open("./models/gmm_model.pkl", "rb") as f:
            gmm_model = pickle.load(f)
        with open("./models/xgboost_model.pkl", "rb") as f:
            xgb_model = pickle.load(f)
            
        return {
            "scaler_features": scaler_features,
            "scaler_gmm": scaler_gmm,
            "gmm_model": gmm_model,
            "xgb_model": xgb_model
        }
    except FileNotFoundError as e:
        st.warning(f"Some model files not found: {str(e)}")
        return None

category_keywords = load_keywords()
job_titles = list(category_keywords.keys())
bert_model = load_bert_model()
models = load_models()

# Text Cleaning Function
def clean_text(text):
    text = str(text).lower()
    text = re.sub(r'[^a-z\s]', '', text)  # Remove special characters
    words = word_tokenize(text)  # Tokenize text
    return words  # Return cleaned word list

# Function to Extract Text from Uploaded Resume
def extract_text_from_file(uploaded_file):
    text = ""
    # If the object has a `.type` attribute, it's likely a Streamlit upload
    if hasattr(uploaded_file, "type"):
        file_type = uploaded_file.type
    else:
        # Fallback for plain file objects: guess from extension
        filename = getattr(uploaded_file, "name", "")
        if filename.endswith(".pdf"):
            file_type = "application/pdf"
        elif filename.endswith(".docx"):
            file_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            file_type = "unknown"

    if file_type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"

    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + "\n"

    else:
        text = "❌ Unsupported file format."

    return text

# Function to Match Resume Against Job Titles Using Keyword Overlap
def match_resume(resume_text):
    resume_words = clean_text(resume_text)

    match_scores = []
    for job_title, keywords in category_keywords.items():
        common_words = set(resume_words) & set(keywords)  # Find overlapping words
        match_score = len(common_words) / len(keywords) if keywords else 0  # Calculate match percentage
        match_scores.append((job_title, match_score))

    # Sort and get top 3 job matches
    match_results = sorted(match_scores, key=lambda x: x[1], reverse=True)[:3]
    return match_results

# Function to compute similarity metrics using imported functions
def compute_similarity_metrics(job_description, resume_text):
    """Calculate similarity metrics between job description and resume text."""
    # Preprocess texts using functions from traditional script
    job_clean = preprocess_text(remove_markup(job_description))
    resume_clean = preprocess_text(remove_markup(resume_text))
    
    # TF-IDF Cosine Similarity - adjusted for single document scenario
    try:
        # For single documents, we need different parameters than the training process
        tfidf_vectorizer = TfidfVectorizer(max_features=5000, ngram_range=(1, 3), min_df=1, max_df=1.0)
        
        # Fit on both texts to ensure we capture all terms
        all_texts = [job_clean, resume_clean]
        tfidf_matrix = tfidf_vectorizer.fit_transform(all_texts)
        
        # Calculate cosine similarity between the two documents
        tfidf_similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    except Exception as e:
        st.warning(f"Error in TF-IDF calculation: {str(e)}")
        tfidf_similarity = 0
    
    # Jaccard Similarity 
    jaccard = jaccard_similarity(job_clean, resume_clean)
    
    # BERT Similarity
    bert_similarity = 0
    if bert_model:
        try:
            # Make sure texts aren't empty
            if job_clean and resume_clean:
                job_embedding = bert_model.encode([job_clean], convert_to_tensor=True)
                resume_embedding = bert_model.encode([resume_clean], convert_to_tensor=True)
                bert_similarity = torch.nn.functional.cosine_similarity(job_embedding, resume_embedding).item()
            else:
                st.warning("Empty text detected for BERT similarity calculation")
        except Exception as e:
            st.warning(f"Error in BERT similarity calculation: {str(e)}")
    
    # N-gram Overlap
    ngram_overlap_score = ngram_overlap(job_clean, resume_clean)
    
    return {
        "tfidf_cosine": tfidf_similarity,
        "jaccard": jaccard,
        "bert_similarity": bert_similarity,
        "ngram_overlap": ngram_overlap_score
    }

# Function to predict match score
def predict_match_score(similarity_metrics):
    """Predict match score using trained models."""
    if not models:
        st.warning("Could not load models. Make sure all model files are available.")
        return None
    
    # Create features as a pandas DataFrame with column names to avoid sklearn warning
    feature_names = ['tfidf_cosine', 'jaccard', 'bert_similarity', 'ngram_overlap']
    features = pd.DataFrame([
        [similarity_metrics["tfidf_cosine"],
         similarity_metrics["jaccard"],
         similarity_metrics["bert_similarity"],
         similarity_metrics["ngram_overlap"]]
    ], columns=feature_names)
    
    try:
        # Scale features using the same scaler used during training
        features_scaled = models["scaler_features"].transform(features)
        
        # Get GMM probabilities
        probabilities = models["gmm_model"].predict_proba(features_scaled)
        gmm_scores = np.dot(probabilities, np.linspace(0, 100, models["gmm_model"].n_components))
        
        # Create DataFrame for GMM scores to avoid warning
        gmm_df = pd.DataFrame(gmm_scores.reshape(-1, 1), columns=['gmm_score'])
        
        # Scale GMM score
        gmm_score_scaled = models["scaler_gmm"].transform(gmm_df).flatten()[0]
        
        # Use XGBoost for final prediction
        xgb_score = models["xgb_model"].predict(features_scaled)[0]
        
        # Return scores individually, converting to native Python float
        return {
            "gmm_score": float(gmm_score_scaled),
            "xgb_score": float(xgb_score)
        }
    except Exception as e:
        st.warning(f"Error predicting match score: {str(e)}")
        return None

# Streamlit UI
st.title("📄 AI-Powered Resume Optimization System")
st.write("Upload your resume for recommendations")

# File Upload
uploaded_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])

# Job Description Input
st.subheader("Optional: Paste Job Description")
job_description = st.text_area("Enter the job description you're applying for", height=200)

if uploaded_file:
    resume_text = extract_text_from_file(uploaded_file)
    
    if resume_text:
        # Add a button to analyze and display results
        analyze_button = st.button("Analyze Resume", type="primary")
        
        if analyze_button:
            # Naive Approach Tab
            st.subheader("Naive Approach: Keyword Matching - Top 3 Job Matches:")
            with st.spinner("Analyzing resume against job categories..."):
                match_results = match_resume(resume_text)
            
            if match_results:
                for job_title, score in match_results:
                    st.write(f"✅ **{job_title}: {score*100:.2f}% match**")
            else:
                st.warning("No matches found based on keyword overlap.")
            
            # Traditional Approach - only if job description is provided
            if job_description:
                st.subheader("Traditional Approach")
                
                with st.spinner("Computing similarity metrics..."):
                    # Compute similarity metrics
                    similarity_metrics = compute_similarity_metrics(job_description, resume_text)
                
                if similarity_metrics:
                    st.write("**Similarity Metrics:**")
                    metrics_df = pd.DataFrame({
                        'Metric': [
                            'TF-IDF Cosine Similarity', 
                            'Jaccard Similarity', 
                            'BERT Semantic Similarity', 
                            'N-gram Overlap'
                        ],
                        'Description': [
                            'Measures content similarity based on term frequency',
                            'Measures word overlap between job and resume',
                            'Captures semantic meaning beyond exact word matches',
                            'Measures phrase overlap between job and resume'
                        ],
                        'Score': [
                            f"{similarity_metrics['tfidf_cosine']:.3f}",
                            f"{similarity_metrics['jaccard']:.3f}",
                            f"{similarity_metrics['bert_similarity']:.3f}",
                            f"{similarity_metrics['ngram_overlap']:.3f}"
                        ]
                    })
                    st.table(metrics_df)
                    
                    # Predict match score
                    match_scores = predict_match_score(similarity_metrics)
                    if match_scores:
                        st.write("**GMM Model Score:**")
                        st.info("Analyzes your resume's similarity features and assigns a score based on clusters of labeled data")
                        gmm_progress_value = min(max(match_scores["gmm_score"] / 100, 0.0), 1.0)
                        st.progress(gmm_progress_value)
                        st.write(f"{match_scores['gmm_score']:.1f}%")
                        
                        st.write("**XGBoost Model Score:**")
                        st.info("Model that predicts match percentage based on the features and labels from training data")
                        xgb_progress_value = min(max(match_scores["xgb_score"] / 100, 0.0), 1.0)
                        st.progress(xgb_progress_value)
                        st.write(f"{match_scores['xgb_score']:.1f}%")
                    else:
                        st.warning("Failed to predict match scores.")
                else:
                    st.warning("Failed to compute similarity metrics.")
            
            # Deep Learning Approach - only if job description is provided
if job_description and process_resume:
    st.subheader("Deep Learning Approach: Llama-Powered Resume Optimization")
    with st.spinner("Optimizing resume using Llama..."):
        try:
            doc = Document(uploaded_file)
        except Exception as e:
            st.error(f"❌ Error opening .docx file: {e}")

        edited_resume_path = process_resume(uploaded_file, job_description)
        
        if edited_resume_path:
            st.success("Resume optimization complete!")

            # Evaluate the optimized resume
            with st.spinner("Evaluating optimized resume..."):
                try:
                    # Save uploaded file to temp location
                    temp_path = "temp_uploaded_resume.docx"
                    with open(temp_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())

                    # Evaluate resume using your DL-based approach
                    deep_results = evaluate_resume(temp_path, edited_resume_path, job_description)

                    st.write("### Deep Learning Evaluation Metrics")
                    st.metric("Keyword Inclusion Score", f"{deep_results['keyword_inclusion_score']*100:.1f}%")
                    st.metric("Semantic Similarity", f"{deep_results['semantic_similarity']:.2f}")

                    os.remove(temp_path)  # Clean up temp file
                except Exception as e:
                    st.error(f"❌ Error during deep evaluation: {e}")
            
            # Download optimized resume
            with open(edited_resume_path, "rb") as file:
                st.download_button(
                    label="Download Optimized Resume",
                    data=file,
                    file_name="optimized_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("Failed to optimize the resume. Please check the input and try again.")
